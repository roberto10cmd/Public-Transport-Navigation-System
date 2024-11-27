import csv
import json
import tkinter as tk
from tkinter import messagebox
from docx import Document
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure

from Model.Repository.NetworkRepository import NetworkRepository
from Model.Repository.UserRepository import UserRepository
from Model.Subject import Subject
from View.EmployeeView import EmployeeView
from Model.Language import Language
import xml.etree.ElementTree as ET
import matplotlib.pyplot as plt
import networkx as nx


class EmployeeController(Subject):

    def __init__(self,empl_view:EmployeeView,user_repo:UserRepository):
        super().__init__()
        self.empl_view=empl_view
        self.user_repo=user_repo
        self.network_repo=NetworkRepository()
        self.add_observer(empl_view)
        self.language=Language()
        self.language.add_observer(empl_view)

        # Set up bindings between view buttons and controller methods
        self.empl_view.LoginButton.configure(command=self.go_to_login_page)
        self.empl_view.add_button.configure(command=self.show_add_gadgets)
        self.empl_view.delete_button.configure(command=self.show_delete_gadgets)
        self.empl_view.update_button.configure(command=self.show_update_gadgets)
        self.empl_view.add_conn_button.configure(command=self.show_add_conn_gadgets)
        self.empl_view.submit_conn_button.configure(command=self.add_connection_method)
        self.empl_view.save_csv.configure(command=self.save_as_csv_method)
        self.empl_view.save_json.configure(command=self.save_as_json_method)
        self.empl_view.save_xml.configure(command=self.save_as_xml_method)
        self.empl_view.save_doc.configure(command=self.save_as_doc_method)
        self.empl_view.delete_conn_button.configure(command=self.show_delete_conn_gadgets)
        self.empl_view.submit__delete_conn_button.configure(command=self.delete_connection_method)
        self.empl_view.submit_add_button.configure(command=self.add_line_method)
        self.empl_view.submit_delete_button.configure(command=self.delete_line_method)
        self.empl_view.submit_update_button_update.configure(command=self.update_line_method)
        self.empl_view.language_combo_box.bind('<<ComboboxSelected>>', self.set_language)
        self.empl_view.show_graph_button.configure(command=self.show_connectivity_chart)

    def set_language(self,event=None):
        print("schimbarea limbii")
        index_of_language=self.empl_view.language_combo_box.current()
        self.notify_observers('change_language',index_of_language)

    def add_line_method(self):
        number = self.empl_view.number_var.get()
        stations_list = self.empl_view.stations_list_var.get()
        line_data=number,stations_list
        self.network_repo.AddLine(line_data)
        self.notify_observers("line_added")

    def delete_line_method(self):
        id = self.empl_view.id_delete_var.get()
        self.network_repo.DeleteLine(id)
        self.notify_observers("line_deleted")


    def update_line_method(self):
        id=self.empl_view.id_update_var.get()
        number=self.empl_view.number_update_var.get()
        stations=self.empl_view.stations_update_var.get()
        self.network_repo.UpdateLine(id, number, stations)
        self.notify_observers("line_updated")

    def delete_connection_method(self):
        id=self.empl_view.id_conn_var.get()
        self.network_repo.delete_connection(id)
        self.notify_observers("connection_deleted")


    def add_connection_method(self):
        start_station_=self.empl_view.start_station_var.get()
        end_station_=self.empl_view.end_station_var.get()
        distance_=self.empl_view.distance_var.get()
        print(start_station_,end_station_,distance_)
        conn_data=end_station_,start_station_ ,distance_
        self.network_repo.addConnection(conn_data)
        self.notify_observers("connection_added")

    def save_as_csv_method(self):
        with open('lines.csv', mode='w', newline='', encoding='utf-8') as file:
            lines_list = self.network_repo.LinesList()
            writer = csv.writer(file)
            writer.writerow(['Line ID', 'Line Number', 'Stations'])
            for line in lines_list:
                writer.writerow([line['line_id'], line['number'], line['stations']])
                messagebox.showinfo("Success!", "Successfully saved as CSV!")
        self.notify_observers("saved_as_csv")


    def save_as_json_method(self):
        lines_list = self.network_repo.LinesList()
        # Ensure to access dictionary keys correctly
        data = [{'line_id': line['line_id'], 'number': line['number'], 'stations': line['stations']} for line in
                lines_list]
        with open('line.json', 'w', encoding='utf-8') as file:
            json.dump(data, file, ensure_ascii=False, indent=4)
        messagebox.showinfo("Success!", "Successfully saved as JSON!")
        self.notify_observers("saved_as_json")


    def save_as_xml_method(self):
        lines_list = self.network_repo.LinesList()  # Correct method to obtain the list
        root = ET.Element('Lines')
        for line in lines_list:
            # Create a Line element with attributes from dictionary keys
            line_element = ET.SubElement(root, 'Line', {'id': str(line['line_id']), 'number': line['number']})
            # Handling the situation where there might be no stations
            stations = line['stations'].split(', ') if line['stations'] else []
            for station in stations:
                ET.SubElement(line_element, 'Station').text = station
        # Convert the XML tree to a string
        xml_str = ET.tostring(root, encoding='unicode')
        # Adding new lines after each Line element for better readability
        formatted_xml_str = xml_str.replace('</Line>', '</Line>\n')
        # Writing the formatted string to a file
        with open('lines.xml', 'w', encoding='utf-8') as file:
            file.write('<?xml version="1.0" encoding="UTF-8"?>\n')  # Add XML declaration at the beginning
            file.write(formatted_xml_str)
        messagebox.showinfo("Success!", "Successfully saved as XML!")
        self.notify_observers("saved_as_xml")


    def save_as_doc_method(self):
        lines_list = self.network_repo.LinesList()

        doc = Document()
        for line in lines_list:
            # Handle station list to ensure it's a string for concatenation
            station_list = line['stations'] if line['stations'] else "No stations"
            doc.add_paragraph(f'Line ID: {line["line_id"]}, Number: {line["number"]}, Stations: {station_list}')
            doc.add_paragraph()  # Adds a blank paragraph for spacing
        doc.save('lines.docx')
        messagebox.showinfo("Success!", "Successfully saved as DOC!")
        self.notify_observers("saved_as_doc")

    def go_to_login_page(self):
        from Controller.LoginController import LoginController
        from View.LoginView import LoginView
        self.empl_view.root.withdraw()
        login_window=tk.Toplevel()
        login_view=LoginView(login_window)
        user_repo=UserRepository()
        login_controller=LoginController(login_view,user_repo)
        login_window.protocol("WM_DELETE_WINDOW", lambda: login_window.destroy())
        self.notify_observers("to_login_page")

    def make_invizible_add(self):
        # Facem widget-urile pentru adăugare invizibile
        self.empl_view.number_label.place_forget()
        self.empl_view.stations_label.place_forget()
        self.empl_view.number_txt_field.place_forget()
        self.empl_view.station_txt_field.place_forget()
        self.empl_view.submit_add_button.place_forget()

    def make_invizible_delete(self):
        self.empl_view.Id_label.place_forget()
        self.empl_view.Id_txt_field.place_forget()
        self.empl_view.submit_delete_button.place_forget()

    def make_invizible_update(self):
        # Facem widget-urile pentru update invizibile
        self.empl_view.number_label_update.place_forget()
        self.empl_view.stations_label_update.place_forget()
        self.empl_view.number_txt_field_update.place_forget()
        self.empl_view.station_txt_field_update.place_forget()
        self.empl_view.id_label_update_update.place_forget()
        self.empl_view.id_txt_update_update.place_forget()
        self.empl_view.submit_update_button_update.place_forget()

    def make_invizible_add_conn(self):
        self.empl_view.start_station_label.place_forget()
        self.empl_view.start_station_field.place_forget()
        self.empl_view.end_station_field.place_forget()
        self.empl_view.end_station_label.place_forget()
        self.empl_view.distance_field.place_forget()
        self.empl_view.distance_label.place_forget()
        self.empl_view.submit_conn_button.place_forget()

    def make_invizible_delete_conn(self):
        self.empl_view.id_coon_label.place_forget()
        self.empl_view.id_coon_field.place_forget()
        self.empl_view.submit__delete_conn_button.place_forget()

    def show_delete_conn_gadgets(self):
        self.make_invizible_add_conn()
        self.empl_view.id_coon_label.place(x=450, y=300)
        self.empl_view.id_coon_label.config(state="normal")
        self.empl_view.id_coon_field.place(x=580, y=300)
        self.empl_view.id_coon_field.config(state="normal")
        self.empl_view.submit__delete_conn_button.place(x=590, y=420)
        self.empl_view.submit__delete_conn_button.config(state="normal")

    def show_add_conn_gadgets(self):
        self.make_invizible_delete_conn()
        self.empl_view.start_station_field.place(x=580, y=340)
        self.empl_view.start_station_field.config(state="normal")

        self.empl_view.end_station_field.place(x=580, y=300)
        self.empl_view.end_station_field.config(state="normal")

        self.empl_view.distance_field.place(x=580, y=380)
        self.empl_view.distance_field.config(state="normal")

        self.empl_view.start_station_label.place(x=450, y=300)
        self.empl_view.start_station_label.config(state="normal")

        self.empl_view.end_station_label.place(x=450, y=340)
        self.empl_view.end_station_label.config(state="normal")

        self.empl_view.distance_label.place(x=450, y=380)
        self.empl_view.distance_label.config(state="normal")

        self.empl_view.submit_conn_button.place(x=600, y=420)
        self.empl_view.submit_conn_button.config(state="normal")

    def show_add_gadgets(self):
        self.make_invizible_delete()
        self.make_invizible_update()

        self.empl_view.number_label.place(x=240, y=60)
        self.empl_view.number_label.config(state="normal")
        self.empl_view.stations_label.place(x=240, y=95)
        self.empl_view.stations_label.config(state="normal")
        self.empl_view.number_txt_field.place(x=320, y=60)
        self.empl_view.number_txt_field.config(state="normal")
        self.empl_view.station_txt_field.place(x=320, y=100)
        self.empl_view.station_txt_field.config(state="normal")
        self.empl_view.submit_add_button.place(x=340, y=140)
        self.empl_view.submit_add_button.config(state="normal")

    def show_delete_gadgets(self):
        self.make_invizible_add()
        self.make_invizible_update()
        self.empl_view.Id_label.place(x=280, y=57)
        self.empl_view.Id_label.config(state="normal")
        self.empl_view.Id_txt_field.place(x=320, y=60)
        self.empl_view.Id_txt_field.config(state="normal")
        self.empl_view.submit_delete_button.place(x=340, y=140)
        self.empl_view.submit_delete_button.config(state="normal")

    def show_update_gadgets(self):
        self.make_invizible_delete()
        self.make_invizible_add()
        self.empl_view.id_label_update_update.place(x=240, y=20)
        self.empl_view.id_label_update_update.config(state="normal")
        self.empl_view.id_txt_update_update.place(x=320, y=20)
        self.empl_view.id_txt_update_update.config(state="normal")
        self.empl_view.number_label_update.place(x=240, y=60)
        self.empl_view.number_label_update.config(state="normal")
        self.empl_view.stations_label_update.place(x=240, y=95)
        self.empl_view.stations_label_update.config(state="normal")
        self.empl_view.number_txt_field_update.place(x=320, y=60)
        self.empl_view.number_txt_field_update.config(state="normal")
        self.empl_view.station_txt_field_update.place(x=320, y=100)
        self.empl_view.station_txt_field_update.config(state="normal")
        self.empl_view.submit_update_button_update.place(x=340, y=140)
        self.empl_view.submit_update_button_update.config(state="normal")

    def show_connectivity_chart(self):
        # Obțineți datele de conectivitate
        connections = self.network_repo.get_all_connections()
        # Crează un grafic pentru rețeaua de stații
        G = nx.DiGraph()

        for conn in connections:
            G.add_edge(conn[0], conn[1], weight=conn[2])

        # Obține lista tuturor nodurilor din graf
        all_nodes = list(G.nodes())

        # Calculează gradul pentru fiecare nod
        degrees = dict(G.degree())

        # Adaugă nodurile cu grad 0 la dicționarul de grade
        for node in all_nodes:
            if node not in degrees:
                degrees[node] = 0

        # Sortează nodurile în ordinea alfabetică
        sorted_nodes = sorted(degrees.keys())

        # Obține gradele pentru nodurile sortate
        connectivity = [degrees[node] for node in sorted_nodes]

        # Crează figura pentru grafic
        fig = Figure(figsize=(10, 5))
        ax = fig.add_subplot(111)
        ax.bar(sorted_nodes, connectivity, color='b')
        ax.set_xlabel('Stations')
        ax.set_ylabel('Connectivity')
        ax.set_title('Station Connectivity in the Network')
        ax.set_xticklabels(sorted_nodes, rotation=45)

        # Creează o nouă fereastră Tkinter
        new_window = tk.Toplevel(self.empl_view.root)
        new_window.title("Station Connectivity Chart")

        # Adaugă graficul la fereastra Tkinter
        canvas = FigureCanvasTkAgg(fig, master=new_window)
        canvas_widget = canvas.get_tk_widget()
        canvas_widget.pack(fill=tk.BOTH, expand=True)
        canvas.draw()

